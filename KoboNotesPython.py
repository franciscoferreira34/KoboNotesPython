# -*- coding: utf-8 -*-
"""
Created on Sun Jul  8 20:20:04 2018

@author: francisco
"""

import pandas as pd
import sqlite3
import warnings
warnings.filterwarnings("ignore")
import os
from docx import Document
import sys

def loadDatabase(file):
    
    conn = sqlite3.connect(file)
    
    ## Bookmark table setup
    
    bookmark = pd.read_sql_query("SELECT * FROM Bookmark", conn)

    #formatar colunas
    bookmark['DateCreated'] = pd.to_datetime(bookmark['DateCreated'])
    bookmark['DateModified'] = pd.to_datetime(bookmark['DateModified'])

    #sort
    bookmark = bookmark.sort_values(['DateCreated','DateModified'], ascending=[1,1])

    # select columns
    bookmark= bookmark[['VolumeID','Text','Annotation','DateCreated','DateModified']]

    bookmark = bookmark.dropna(subset=['Text'])

    uniqueVolumeIDfromBookmark = bookmark.VolumeID.unique()

    #conn.text_factory = bytes
    content = pd.read_sql_query("SELECT ContentID,ContentType,MimeType,Title,Attribution FROM content", conn)

    #filter content by books with highlights and/or annotations
    content = content[content.ContentID.isin(uniqueVolumeIDfromBookmark)]

    #sort
    content = content.sort_values(['Title'], ascending=[1])
    
    content = content.drop_duplicates(subset=['ContentID'])
    
    conn.close()
    
    return bookmark,content
    
def createUnifiedDatabase(bookmark,content):
    
    bookmark['Title']='None'
        
    for i in range(len(content)):
        bookmark['Title'][bookmark.VolumeID==content.iloc[i].ContentID]=content.iloc[i].Title
        
    bookmark = bookmark.sort_values(['Title','DateCreated'], ascending=[1,1])
    
    return bookmark
    
def getBooks(file):
    
    bookmark,content = loadDatabase(file)
    
    unique_titles = content.Title.unique()
    
    print('\n'.join(unique_titles))
    
def BookToTXT(file,book):
    
    bookmark,content = loadDatabase(file)
    
    data = createUnifiedDatabase(bookmark,content)
    
    data = data[data.Title==book]
    
    if(len(data)==0):
        print("Book not available. Try one of the following books:\n")
        print(getBooks(file))
    else:
        _saveTXTFile(data['Text'][data.Title==book].get_values(),data['Annotation'][data.Title==book].get_values(),book+".txt")

def saveBooksToTXT(file,folder):
    
    bookmark,content = loadDatabase(file)
    
    data = createUnifiedDatabase(bookmark,content)

    
    books = data.Title.unique()
    
    if not os.path.exists(folder):
        os.makedirs(folder)
    
    for book in books:
        
        _saveTXTFile(data['Text'][data.Title==book].get_values(),data['Annotation'][data.Title==book].get_values(),os.path.join(folder,book+".txt"))

def _saveTXTFile(arrayText,arrayAnnotation,nameFile):
    with open(nameFile, "w") as f:
        for text,annotation in zip(arrayText,arrayAnnotation):
            if annotation!=None and annotation!="":
                f.write(text+" Annotation: "+annotation +"\n\n")
            else:
                f.write(text+"\n\n")
                
def saveBooksToWord(file,out_path):
    
    document = Document()
    
    bookmark,content = loadDatabase(file)
    
    data = createUnifiedDatabase(bookmark,content)
    
    books = data.Title.unique()
    
    if not os.path.exists(os.path.dirname(out_path)) and os.path.dirname(out_path)!="":
        os.makedirs(os.path.dirname(out_path))
    
    for book in books:
        
        arrayText = data['Text'][data.Title==book].get_values()
        arrayAnnotation = data['Annotation'][data.Title==book].get_values()
        
        document.add_heading(book, 1)
        
        p = document.add_paragraph()
        
        for text,annotation in zip(arrayText,arrayAnnotation):
            while(text[0]=="." or text[0]==" " or text[0]=="," or text[0]==":"):
                text=text[1:]
            text = text.replace('\n','')
            p = document.add_paragraph(text)
            if annotation!=None and annotation!="":
                p.add_run(' Annotation: ').bold = True
                p.add_run(annotation)
        
        document.add_page_break()
        
    document.save(out_path)


if __name__ == "__main__":
    
    if sys.argv[1]=='toDOCX':
        
        if(len(sys.argv)>=4):
        
            saveBooksToWord(sys.argv[2],sys.argv[3])
            print('DOCX created')
        
        else:
            
            print("Please provide file .sqlite path and out path for .docx")
    
    else:
        print('Command not recognized')